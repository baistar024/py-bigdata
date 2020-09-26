import docx, openpyxl, re, os
from win32com import client as wc



def doc2docx(files):
    word = wc.Dispatch("Word.Application")
    for file in files:
        doc = word.Documents.Open(file)
        doc.SaveAs("{}x".format(file), 12)
        doc.Close()
    word.Quit()

def doc2xls(worddocs, wbexcel):
    i = 0
    datapatter = r'2020.\s*[0|1]?\d.\s*[0|1|2|3]?\d'
    dormpatter = r'1[4|5|7|8]号'
    doc = docx.Document(worddoc)
    texts = []
    for para in doc.paragraphs:
        if "日期" in para.text and "号" in para.text:
            texts.append(para.text)

    wb = openpyxl.load_workbook(wbexcel)
    ws = wb.active

    tables = doc.tables
    for table in tables:
        date = getDate(texts[0])
        for row in table.rows:
            row_content = []
            if row.cells[1].text =="艺术":
                row_content.append(date)
                row_content.append(getDrom(texts[i]))
                for cell in row.cells:
                    row_content.append(cell.text)
                ws.append(row_content)
        i += 1
    wb.save("score.xlsx")
    wb.close()
    doc.Close()
def getDate(text):
    datapatter = r'2020.\s*[0|1]?\d.\s*[0|1|2|3]?\d'
    return re.search(datapatter,text).group()
def getDrom(text):
    dormpatter = r'1[4|5|7|8]号'
    return re.search(dormpatter,text).group()

def main():
    path = os.path.join(os.getcwd(), "record")
    files = []
    for file in os.listdir(os.path.join(path)):
        if file.endswith(".doc"):
            files.append(os.path.join(path, file))
    doc2docx(files)

main()